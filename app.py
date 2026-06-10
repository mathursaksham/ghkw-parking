import logging
import os
import random
import re
import smtplib
import subprocess
import zipfile
from datetime import datetime
from email.mime.application import MIMEApplication
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO

import pandas as pd
import streamlit as st
from docx import Document

# --- 1. PAGE CONFIGURATION ---
# This MUST be the first Streamlit command and flush against the left margin.
st.set_page_config(page_title="GHKW Portal", page_icon="🔒", layout="wide")

# --- 2. LOGGING CONFIGURATION ---
logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# --- 3. CONFIGURATION & ENV VARIABLES ---
EXCEL_FILE = "data.xlsx"  # Path to your Excel file
TEMPLATE_FILE = "template.docx"  # Path to your Word template

SUPERADMIN_EMAIL = "ghkwparkingallotments@gmail.com"

SMTP_SERVER = os.getenv("SMTP_SERVER", "smtp.gmail.com")
SMTP_PORT = int(os.getenv("SMTP_PORT", 587))
SENDER_EMAIL = os.getenv("SENDER_EMAIL", "ghkwparkingallotments@gmail.com")
SENDER_PASSWORD = os.getenv("SENDER_PASSWORD", "mjqy ohdf beeg vngh")

# Set max log size (1 MB = 1 * 1024 * 1024 bytes)
MAX_LOG_SIZE_BYTES = 1 * 1024 * 1024 


# --- 4. AUTO-EMAIL LOGS UTILITY ---
def check_and_email_logs():
    """Checks log size. If over limit, emails it to admin and clears the log."""
    log_file_path = "app.log"
    
    # Check if file exists and if its size exceeds the limit
    if os.path.exists(log_file_path) and os.path.getsize(log_file_path) > MAX_LOG_SIZE_BYTES:
        try:
            msg = MIMEMultipart()
            msg["From"] = SENDER_EMAIL
            msg["To"] = SUPERADMIN_EMAIL
            msg["Subject"] = f"Automated Log Backup: {datetime.now().strftime('%Y-%m-%d')}"

            body = "The application log has reached the size limit. The backup is attached."
            msg.attach(MIMEText(body, "plain"))

            # Attach the log file
            with open(log_file_path, "rb") as f:
                attach = MIMEApplication(f.read(), _subtype="txt")
                attach.add_header(
                    "Content-Disposition", 
                    "attachment", 
                    filename=f"ghkw_audit_log_{datetime.now().strftime('%Y%m%d')}.txt"
                )
                msg.attach(attach)

            # Send Email
            server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
            server.starttls()
            server.login(SENDER_EMAIL, SENDER_PASSWORD)
            server.send_message(msg)
            server.quit()

            # If email succeeds, completely clear the log file to start fresh
            with open(log_file_path, "w") as f:
                f.truncate(0)
                
            # Log that a backup just happened
            logging.info("SYSTEM | Log file reached capacity. Automatically emailed to admin and wiped clean.")

        except Exception as e:
            # We fail silently here so the app doesn't crash for regular users if SMTP hits a limit
            print(f"Failed to auto-email logs: {e}")

# Run the check every time the script executes
check_and_email_logs()


# --- 5. AUTHENTICATION FUNCTIONS ---
def send_otp_email(receiver_email, otp):
    """Sends a 6-digit OTP to the user's email."""
    msg = MIMEMultipart()
    msg["From"] = SENDER_EMAIL
    msg["To"] = receiver_email
    msg["Subject"] = "Your Verification Code - GHKW Parking Allotments"

    body = f"""
    <h3>GHKW Parking Allotments Portal</h3>
    <p>You requested access to the portal. Please use the verification code below:</p>
    <h2 style="color: #4F8BF9; letter-spacing: 2px;">{otp}</h2>
    <p>This code will expire shortly. If you did not request this, please ignore this email.</p>
    """
    msg.attach(MIMEText(body, "html"))

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        # 📝 AUDIT LOG: Silently log the real technical error for the admin
        logging.error(f"SMTP Email Failure (Likely Limit Exceeded): {e}")
        
        # Show the friendly message to the user
        st.error("The Site is experiencing high volume please retry in sometime.")
        return False


def render_login_page():
    """Renders the login and OTP validation UI."""
    # Center-align login container
    _, auth_col, _ = st.columns([1, 2, 1])

    with auth_col:
        st.markdown(
            "<h2 style='text-align: center;'>🔒 GHKW Portal Authentication</h2>",
            unsafe_allow_html=True,
        )
        st.write("Please authenticate with your registered email to gain access.")

        # Step 1: Input Email
        email_input = st.text_input(
            "Enter Email ID",
            placeholder="user@example.com",
            disabled=st.session_state.otp_sent,
        ).strip()

        # Step 1 Button: Request OTP
        if not st.session_state.otp_sent:
            if st.button("Send OTP", use_container_width=True):
                if not re.match(r"[^@]+@[^@]+\.[^@]+", email_input):
                    st.error("Please enter a valid email address.")
                else:
                    generated_otp = str(random.randint(100000, 999999))
                    with st.spinner("Sending security code..."):
                        if send_otp_email(email_input, generated_otp):
                            st.session_state.otp_sent = True
                            st.session_state.target_email = email_input
                            st.session_state.generated_otp = generated_otp
                            st.success(f"OTP successfully sent to {email_input}")
                            st.rerun()

        # Step 2: Input OTP
        if st.session_state.otp_sent:
            st.info(f"A code has been sent to **{st.session_state.target_email}**")
            otp_input = st.text_input("Enter 6-Digit OTP", placeholder="******").strip()

            col_verify, col_reset = st.columns(2)
            with col_verify:
                if st.button("Verify & Login", use_container_width=True):
                    if otp_input == st.session_state.generated_otp:
                        st.session_state.authenticated = True
                        logging.info(f"USER_LOGIN | Email: {st.session_state.target_email} | Status: SUCCESS")
                        st.success("Access Granted!")
                        st.rerun()
                    else:
                        st.error("Invalid verification code. Please try again.")

            with col_reset:
                if st.button("Change Email / Resend", use_container_width=True):
                    st.session_state.otp_sent = False
                    st.session_state.generated_otp = None
                    st.rerun()


# --- 6. INITIALIZE SESSION STATE ---
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "otp_sent" not in st.session_state:
    st.session_state.otp_sent = False
if "generated_otp" not in st.session_state:
    st.session_state.generated_otp = None
if "target_email" not in st.session_state:
    st.session_state.target_email = ""


# --- 7. ROUTE USER BASED ON AUTH STATUS ---
if not st.session_state.authenticated:
    render_login_page()
else:
    # --- CORE PARKING APPLICATION ---

    @st.cache_data
    def load_data():
        """Load and clean Excel data."""
        if not os.path.exists(EXCEL_FILE):
            st.error(f"Excel file '{EXCEL_FILE}' not found!")
            return pd.DataFrame()
        return pd.read_excel(EXCEL_FILE, dtype=str)

    def generate_pdf_bytes(row_data, flat_number):
        """Populates Word template and converts to PDF using LibreOffice."""
        if not os.path.exists(TEMPLATE_FILE):
            st.error(f"Template file '{TEMPLATE_FILE}' not found!")
            return None

        os.makedirs("/tmp", exist_ok=True)
        doc = Document(TEMPLATE_FILE)
        row_data["Date"] = datetime.now().strftime("%B %d, %Y")

        for paragraph in doc.paragraphs:
            for key, value in row_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in row_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(placeholder, str(value))

        tmp_docx = f"/tmp/Letter_Flat_{flat_number}.docx"
        pdf_path = f"/tmp/Letter_Flat_{flat_number}.pdf"
        doc.save(tmp_docx)

        try:
            libreoffice_cmd = "libreoffice"
            cmd = [
                libreoffice_cmd,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                "/tmp",
                tmp_docx,
            ]
            subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                env={"HOME": "/tmp"},
            )

            if not os.path.exists(pdf_path):
                st.error(f"LibreOffice conversion finished but no PDF found for Flat {flat_number}.")
                return None

            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

            if os.path.exists(tmp_docx):
                os.remove(tmp_docx)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            return pdf_bytes
        except Exception as e:
            st.error(f"Error converting Flat {flat_number} via LibreOffice: {e}")
            return None

    # --- SIDEBAR (Includes Reload Data & Logout) ---
    with st.sidebar:
        st.write(f"👤 Logged in as: **{st.session_state.target_email}**")
        
        st.write("---")
        # Clear Cache Button
        if st.button("🔄 Reload Excel Data", use_container_width=True):
            st.cache_data.clear()
            st.success("Cache cleared! Fetching latest Excel data...")
            st.rerun()

        # Sign Out Button
        if st.button("Sign Out", use_container_width=True):
            logging.info(f"USER_LOGOUT | Email: {st.session_state.target_email}")
            st.session_state.authenticated = False
            st.session_state.otp_sent = False
            st.session_state.generated_otp = None
            st.rerun()

    # --- MAIN INTERFACE ---
    st.title("🏢 GHKW Parking Allotments")
    df = load_data()

    if not df.empty:
        if "Flat Number" not in df.columns:
            st.error("Excel sheet must contain a 'Flat Number' column.")
        else:
            if "selected_flats_tracker" not in st.session_state:
                st.session_state.selected_flats_tracker = set()

            col1, col2 = st.columns([3, 2])

            with col1:
                st.header("🔍 Search & Select Flats")
                search_query = st.text_input(
                    "Search by Flat / Block (e.g., 'C5', 'A', 'A1-003'):",
                    placeholder="Type to filter rows...",
                ).strip()

                if search_query:
                    filtered_df = df[
                        df["Flat Number"].str.contains(search_query, case=False, na=False)
                    ].copy()
                else:
                    filtered_df = df.copy()

                filtered_df.insert(
                    0,
                    "Select",
                    filtered_df["Flat Number"].apply(lambda x: x in st.session_state.selected_flats_tracker),
                )

                btn_col1, btn_col2 = st.columns(2)
                with btn_col1:
                    if st.button("✅ Check All Filtered Rows"):
                        for f_num in filtered_df["Flat Number"]:
                            st.session_state.selected_flats_tracker.add(f_num)
                        st.rerun()
                with btn_col2:
                    if st.button("❌ Uncheck All Filtered Rows"):
                        for f_num in filtered_df["Flat Number"]:
                            st.session_state.selected_flats_tracker.discard(f_num)
                        st.rerun()

                edited_df = st.data_editor(
                    filtered_df,
                    hide_index=True,
                    column_config={
                        "Select": st.column_config.CheckboxColumn(
                            "Select",
                            help="Check to include this flat",
                            default=False,
                        )
                    },
                    disabled=[col for col in filtered_df.columns if col != "Select"],
                    width="stretch",
                    key="flat_data_editor",
                )

                if edited_df is not None:
                    for _, row in edited_df.iterrows():
                        f_num = row["Flat Number"]
                        if row["Select"]:
                            st.session_state.selected_flats_tracker.add(f_num)
                        else:
                            st.session_state.selected_flats_tracker.discard(f_num)

                current_selections = list(st.session_state.selected_flats_tracker)
                selected_rows = df[df["Flat Number"].isin(current_selections)]

                st.write(f"📂 **Total unique flats selected overall:** {len(current_selections)}")
                if len(current_selections) > 0:
                    with st.expander("See selected list"):
                        st.write(", ".join(sorted(current_selections)))

                if st.button("Generate Letters for Chosen Rows"):
                    if not current_selections:
                        st.warning("Please select or search-check at least one flat checkbox above.")
                    elif len(current_selections) == 1:
                        flat = current_selections[0]
                        row = selected_rows[selected_rows["Flat Number"] == flat].iloc[0].to_dict()

                        with st.spinner(f"Processing PDF for Flat {flat}..."):
                            pdf_data = generate_pdf_bytes(row, flat)
                            if pdf_data:
                                logging.info(f"ACTION | User: {st.session_state.target_email} | Generated single PDF for Flat: {flat}")
                                st.success(f"PDF for Flat {flat} ready!")
                                st.download_button(
                                    label="⬇️ Save PDF to Desktop",
                                    data=pdf_data,
                                    file_name=f"Letter_Flat_{flat}.pdf",
                                    mime="application/pdf",
                                )
                    else:
                        with st.spinner("Processing chosen layout configurations..."):
                            zip_buffer = BytesIO()
                            progress_bar = st.progress(0)
                            total = len(current_selections)
                            processed_count = 0

                            with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                                for index, flat in enumerate(current_selections):
                                    row = selected_rows[selected_rows["Flat Number"] == flat].iloc[0].to_dict()
                                    pdf_data = generate_pdf_bytes(row, flat)
                                    if pdf_data:
                                        zip_file.writestr(f"Letter_Flat_{flat}.pdf", pdf_data)
                                        processed_count += 1
                                    progress_bar.progress((index + 1) / total)

                            if processed_count > 0:
                                logging.info(f"ACTION | User: {st.session_state.target_email} | Generated ZIP archive for {processed_count} flats: {current_selections}")
                                st.success(f"Successfully packaged {processed_count} letters!")
                                st.download_button(
                                    label="⬇️ Download Selected PDFs (ZIP)",
                                    data=zip_buffer.getvalue(),
                                    file_name="Selected_Parking_Letters.zip",
                                    mime="application/zip",
                                )
                            else:
                                st.error("Failed to generate PDFs for the chosen flats.")

            with col2:
                st.header("🚀 Bulk Parking Print")
                st.write("Compile all letters with assigned parking spaces into a ZIP file.")

                if st.button("Prepare All Parkings"):
                    if "Parking" in df.columns:
                        parking_df = df[df["Parking"].notna() & (df["Parking"] != "")]
                    else:
                        parking_df = df

                    if parking_df.empty:
                        st.warning("No records found with parking details.")
                    else:
                        zip_buffer = BytesIO()
                        progress_bar = st.progress(0)
                        total = len(parking_df)

                        with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
                            for index, (_, row) in enumerate(parking_df.iterrows()):
                                flat_num = row["Flat Number"]
                                pdf_data = generate_pdf_bytes(row.to_dict(), flat_num)
                                if pdf_data:
                                    zip_file.writestr(f"Letter_Flat_{flat_num}.pdf", pdf_data)
                                progress_bar.progress((index + 1) / total)

                        logging.info(f"ACTION | User: {st.session_state.target_email} | Generated BULK ZIP for all {len(parking_df)} assigned parking records")
                        st.success("ZIP package generated successfully!")
                        st.download_button(
                            label="⬇️ Download All PDFs (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name="All_Parking_Letters.zip",
                            mime="application/zip",
                        )

    # --- 8. SECURE COMPLIANCE VIEWER PANEL FOR SUPERADMIN ---
    if st.session_state.target_email == SUPERADMIN_EMAIL:
        st.write("---")
        st.subheader("📋 System Audit Logs (Admin Only)")

        if os.path.exists("app.log"):
            with open("app.log", "r") as f:
                log_content = f.read()

            st.text_area("Live App Activity History", value=log_content, height=250, disabled=True)

            st.download_button(
                label="⬇️ Download Security Log File",
                data=log_content,
                file_name=f"ghkw_audit_log_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
            )
        else:
            st.info("System initializing. Logs are clean.")