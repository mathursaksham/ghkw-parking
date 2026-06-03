import logging
import os
import random
import re
import smtplib
import subprocess
import zipfile
from datetime import datetime
from email.mime.multipart import MIMEMultipart
from email.mime.text import MIMEText
from io import BytesIO
import pandas as pd
import streamlit as st
from docx import Document

# --- 1. LOGGING CONFIGURATION ---
# This automatically sets up or appends to 'app.log' in your root directory
logging.basicConfig(
    filename="app.log",
    level=logging.INFO,
    format="%(asctime)s - %(levelname)s - %(message)s",
    datefmt="%Y-%m-%d %H:%M:%S",
)

# --- 2. CONFIGURATION & SECRETS ---
EXCEL_FILE = "data.xlsx"  # Path to your Excel file
TEMPLATE_FILE = "template.docx"  # Path to your Word template

# Define who can view the raw log file directly inside the app interface
SUPERADMIN_EMAIL = "superadmin@ghkw.com"

# Fetch email configurations securely from st.secrets
try:
    SMTP_SERVER = st.secrets["SMTP_SERVER"]
    SMTP_PORT = int(st.secrets["SMTP_PORT"])
    SENDER_EMAIL = os.getenv("SENDER_EMAIL", "ghkwparkingallotments@gmail.com")
    SENDER_PASSWORD = os.getenv("SENDER_PASSWORD", "mjqy ohdf beeg vngh")  # Use App Password for Gmail

except Exception:
    st.error(
        "Missing secret configurations! Please check your `.streamlit/secrets.toml` file."
    )
    st.stop()


# --- 3. AUTHENTICATION FUNCTIONS ---
def send_otp_email(receiver_email, otp):
    """Sends a 6-digit OTP to the user's email."""
    msg = MIMEMultipart()
    msg["From"] = SENDER_EMAIL
    msg["To"] = receiver_email
    msg["Subject"] = "Your Verification Code - GHKW Parking Allotments"

    body = f"""
    <h3>GHKW Parking Allotments Portal</h3>
    <p>You requested access to the portal. Please use the verification code below:</p>
    <h2 style="color: #4F8BF9; letter-spacing: 2px;">{otp}</h2>
    <p>This code will expire shortly. If you did not request this, please ignore this email.</p>
    """
    msg.attach(MIMEText(body, "html"))

    try:
        server = smtplib.SMTP(SMTP_SERVER, SMTP_PORT)
        server.starttls()
        server.login(SENDER_EMAIL, SENDER_PASSWORD)
        server.send_message(msg)
        server.quit()
        return True
    except Exception as e:
        st.sidebar.error(f"Failed to send email: {e}")
        return False


def render_login_page():
    """Renders the login and OTP validation UI."""
    st.set_page_config(page_title="GHKW Portal Login", page_icon="🔒", layout="centered")

    _, auth_col, _ = st.columns([1, 2, 1])

    with auth_col:
        st.markdown(
            "<h2 style='text-align: center;'>🔒 Portal Authentication</h2>",
            unsafe_allow_html=True,
        )
        st.write("Please authenticate with your registered email to gain access.")

        # Step 1: Input Email
        email_input = st.text_input(
            "Enter Email ID",
            placeholder="user@example.com",
            disabled=st.session_state.otp_sent,
        ).strip()

        # Step 1 Button: Request OTP
        if not st.session_state.otp_sent:
            if st.button("Send OTP", use_container_width=True):
                if not re.match(r"[^@]+@[^@]+\.[^@]+", email_input):
                    st.error("Please enter a valid email address.")
                else:
                    generated_otp = str(random.randint(100000, 999999))
                    with st.spinner("Sending security code..."):
                        if send_otp_email(email_input, generated_otp):
                            st.session_state.otp_sent = True
                            st.session_state.target_email = email_input
                            st.session_state.generated_otp = generated_otp
                            st.success(f"OTP successfully sent to {email_input}")
                            st.rerun()

        # Step 2: Input OTP (Visible only after sending OTP)
        if st.session_state.otp_sent:
            st.info(f"A code has been sent to **{st.session_state.target_email}**")
            otp_input = st.text_input(
                "Enter 6-Digit OTP", placeholder="******"
            ).strip()

            col_verify, col_reset = st.columns(2)
            with col_verify:
                if st.button("Verify & Login", use_container_width=True):
                    if otp_input == st.session_state.generated_otp:
                        st.session_state.authenticated = True

                        # AUDIT LOG: Successful entry
                        logging.info(
                            f"USER_LOGIN | Email: {st.session_state.target_email} | Status: SUCCESS"
                        )

                        st.success("Access Granted!")
                        st.rerun()
                    else:
                        st.error("Invalid verification code. Please try again.")

            with col_reset:
                if st.button("Change Email / Resend", use_container_width=True):
                    st.session_state.otp_sent = False
                    st.session_state.generated_otp = None
                    st.rerun()


# --- 4. INITIALIZE AUTH SESSION STATE ---
if "authenticated" not in st.session_state:
    st.session_state.authenticated = False
if "otp_sent" not in st.session_state:
    st.session_state.otp_sent = False
if "generated_otp" not in st.session_state:
    st.session_state.generated_otp = None
if "target_email" not in st.session_state:
    st.session_state.target_email = ""

# --- 5. ROUTE USER BASED ON AUTH STATUS ---
if not st.session_state.authenticated:
    render_login_page()
else:
    # --- CORE PARKING APPLICATION (Triggers only when Authenticated) ---

    @st.cache_data
    def load_data():
        """Load and clean Excel data."""
        if not os.path.exists(EXCEL_FILE):
            st.error(f"Excel file '{EXCEL_FILE}' not found!")
            return pd.DataFrame()
        return pd.read_excel(EXCEL_FILE, dtype=str)

    def generate_pdf_bytes(row_data, flat_number):
        """Populates Word template and converts to PDF using LibreOffice."""
        if not os.path.exists(TEMPLATE_FILE):
            st.error(f"Template file '{TEMPLATE_FILE}' not found!")
            return None

        os.makedirs("/tmp", exist_ok=True)
        doc = Document(TEMPLATE_FILE)
        row_data["Date"] = datetime.now().strftime("%B %d, %Y")

        # Replace placeholders in paragraphs
        for paragraph in doc.paragraphs:
            for key, value in row_data.items():
                placeholder = f"{{{{{key}}}}}"
                if placeholder in paragraph.text:
                    for run in paragraph.runs:
                        if placeholder in run.text:
                            run.text = run.text.replace(placeholder, str(value))

        # Replace placeholders in tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for key, value in row_data.items():
                            placeholder = f"{{{{{key}}}}}"
                            if placeholder in paragraph.text:
                                paragraph.text = paragraph.text.replace(
                                    placeholder, str(value)
                                )

        tmp_docx = f"/tmp/Letter_Flat_{flat_number}.docx"
        pdf_path = f"/tmp/Letter_Flat_{flat_number}.pdf"
        doc.save(tmp_docx)

        try:
            libreoffice_cmd = "libreoffice"
            cmd = [
                libreoffice_cmd,
                "--headless",
                "--convert-to",
                "pdf",
                "--outdir",
                "/tmp",
                tmp_docx,
            ]
            subprocess.run(
                cmd,
                stdout=subprocess.PIPE,
                stderr=subprocess.PIPE,
                check=True,
                env={"HOME": "/tmp"},
            )

            if not os.path.exists(pdf_path):
                st.error(
                    f"LibreOffice conversion finished but no PDF found for Flat {flat_number}."
                )
                return None

            with open(pdf_path, "rb") as f:
                pdf_bytes = f.read()

            if os.path.exists(tmp_docx):
                os.remove(tmp_docx)
            if os.path.exists(pdf_path):
                os.remove(pdf_path)

            return pdf_bytes
        except Exception as e:
            st.error(
                f"Error converting Flat {flat_number} via LibreOffice: {e}"
            )
            return None

    # --- STREAMLIT FRONTEND ---
    st.set_page_config(layout="wide")

    # Add Logout Button to Sidebar
    with st.sidebar:
        st.write(f"👤 Logged in as: **{st.session_state.target_email}**")
        if st.button("Sign Out"):
            # AUDIT LOG: Explicit Logout Tracker
            logging.info(f"USER_LOGOUT | Email: {st.session_state.target_email}")

            st.session_state.authenticated = False
            st.session_state.otp_sent = False
            st.session_state.generated_otp = None
            st.rerun()

    st.title("🏢 GHKW Parking Allotments")
    df = load_data()

    if not df.empty:
        if "Flat Number" not in df.columns:
            st.error("Excel sheet must contain a 'Flat Number' column.")
        else:
            if "selected_flats_tracker" not in st.session_state:
                st.session_state.selected_flats_tracker = set()

            col1, col2 = st.columns([3, 2])

            with col1:
                st.header("🔍 Search & Select Flats")
                search_query = st.text_input(
                    "Search by Flat / Block (e.g., 'C5', 'A', 'A1-003'):",
                    placeholder="Type to filter rows...",
                ).strip()

                if search_query:
                    filtered_df = df[
                        df["Flat Number"].str.contains(
                            search_query, case=False, na=False
                        )
                    ].copy()
                else:
                    filtered_df = df.copy()

                filtered_df.insert(
                    0,
                    "Select",
                    filtered_df["Flat Number"].apply(
                        lambda x: x in st.session_state.selected_flats_tracker
                    ),
                )

                btn_col1, btn_col2 = st.columns(2)
                with btn_col1:
                    if st.button("✅ Check All Filtered Rows"):
                        for f_num in filtered_df["Flat Number"]:
                            st.session_state.selected_flats_tracker.add(f_num)
                        st.rerun()
                with btn_col2:
                    if st.button("❌ Uncheck All Filtered Rows"):
                        for f_num in filtered_df["Flat Number"]:
                            st.session_state.selected_flats_tracker.discard(
                                f_num
                            )
                        st.rerun()

                edited_df = st.data_editor(
                    filtered_df,
                    hide_index=True,
                    column_config={
                        "Select": st.column_config.CheckboxColumn(
                            "Select",
                            help="Check to include this flat",
                            default=False,
                        )
                    },
                    disabled=[
                        col for col in filtered_df.columns if col != "Select"
                    ],
                    width="stretch",
                    key="flat_data_editor",
                )

                if edited_df is not None:
                    for _, row in edited_df.iterrows():
                        f_num = row["Flat Number"]
                        if row["Select"]:
                            st.session_state.selected_flats_tracker.add(f_num)
                        else:
                            st.session_state.selected_flats_tracker.discard(
                                f_num
                            )

                current_selections = list(
                    st.session_state.selected_flats_tracker
                )
                selected_rows = df[df["Flat Number"].isin(current_selections)]

                st.write(
                    f"📂 **Total unique flats selected overall:** {len(current_selections)}"
                )
                if len(current_selections) > 0:
                    with st.expander("See selected list"):
                        st.write(", ".join(sorted(current_selections)))

                if st.button("Generate Letters for Chosen Rows"):
                    if not current_selections:
                        st.warning(
                            "Please select or search-check at least one flat checkbox above."
                        )

                    elif len(current_selections) == 1:
                        flat = current_selections[0]
                        row = (
                            selected_rows[selected_rows["Flat Number"] == flat]
                            .iloc[0]
                            .to_dict()
                        )

                        with st.spinner(f"Processing PDF for Flat {flat}..."):
                            pdf_data = generate_pdf_bytes(row, flat)
                            if pdf_data:
                                # AUDIT LOG: Single PDF print action
                                logging.info(
                                    f"ACTION | User: {st.session_state.target_email} | Generated single PDF for Flat: {flat}"
                                )

                                st.success(f"PDF for Flat {flat} ready!")
                                st.download_button(
                                    label="⬇️ Save PDF to Desktop",
                                    data=pdf_data,
                                    file_name=f"Letter_Flat_{flat}.pdf",
                                    mime="application/pdf",
                                )

                    else:
                        with st.spinner(
                            "Processing chosen layout configurations..."
                        ):
                            zip_buffer = BytesIO()
                            progress_bar = st.progress(0)
                            total = len(current_selections)
                            processed_count = 0

                            with zipfile.ZipFile(
                                zip_buffer, "w", zipfile.ZIP_DEFLATED
                            ) as zip_file:
                                for index, flat in enumerate(
                                    current_selections
                                ):
                                    row = (
                                        selected_rows[
                                            selected_rows["Flat Number"] == flat
                                        ]
                                        .iloc[0]
                                        .to_dict()
                                    )
                                    pdf_data = generate_pdf_bytes(row, flat)
                                    if pdf_data:
                                        zip_file.writestr(
                                            f"Letter_Flat_{flat}.pdf", pdf_data
                                        )
                                        processed_count += 1

                                    progress_bar.progress((index + 1) / total)

                            if processed_count > 0:
                                # AUDIT LOG: Multiple PDF selection print action
                                logging.info(
                                    f"ACTION | User: {st.session_state.target_email} | Generated ZIP for {processed_count} flats: {current_selections}"
                                )

                                st.success(
                                    f"Successfully packaged {processed_count} letters!"
                                )
                                st.download_button(
                                    label="⬇️ Download Selected PDFs (ZIP)",
                                    data=zip_buffer.getvalue(),
                                    file_name="Selected_Parking_Letters.zip",
                                    mime="application/zip",
                                )
                            else:
                                st.error(
                                    "Failed to generate PDFs for the chosen flats."
                                )

            with col2:
                st.header("🚀 Bulk Parking Print")
                st.write(
                    "Compile all letters with assigned parking spaces into a ZIP file."
                )

                if st.button("Prepare All Parkings"):
                    if "Parking" in df.columns:
                        parking_df = df[
                            df["Parking"].notna() & (df["Parking"] != "")
                        ]
                    else:
                        parking_df = df

                    if parking_df.empty:
                        st.warning("No records found with parking details.")
                    else:
                        zip_buffer = BytesIO()
                        progress_bar = st.progress(0)
                        total = len(parking_df)

                        with zipfile.ZipFile(
                            zip_buffer, "w", zipfile.ZIP_DEFLATED
                        ) as zip_file:
                            for index, (_, row) in enumerate(
                                parking_df.iterrows()
                            ):
                                flat_num = row["Flat Number"]
                                pdf_data = generate_pdf_bytes(
                                    row.to_dict(), flat_num
                                )

                                if pdf_data:
                                    zip_file.writestr(
                                        f"Letter_Flat_{flat_num}.pdf", pdf_data
                                    )

                                progress_bar.progress((index + 1) / total)

                        # AUDIT LOG: Full Database Export Triggered
                        logging.info(
                            f"ACTION | User: {st.session_state.target_email} | Generated BULK ZIP for all {len(parking_df)} parking records"
                        )

                        st.success("ZIP package generated successfully!")
                        st.download_button(
                            label="⬇️ Download All PDFs (ZIP)",
                            data=zip_buffer.getvalue(),
                            file_name="All_Parking_Letters.zip",
                            mime="application/zip",
                        )

    # --- 6. SECURE COMPLIANCE VIEWER FOR SUPERADMIN ---
    if st.session_state.target_email == SUPERADMIN_EMAIL:
        st.write("---")
        st.subheader("📋 System Audit Logs (Admin Only)")

        if os.path.exists("app.log"):
            with open("app.log", "r") as f:
                log_content = f.read()

            st.text_area(
                "Live App Activity History",
                value=log_content,
                height=250,
                disabled=True,
            )

            st.download_button(
                label="⬇️ Download Security Log File",
                data=log_content,
                file_name=f"ghkw_audit_log_{datetime.now().strftime('%Y%m%d')}.txt",
                mime="text/plain",
            )
        else:
            st.info("System initializing. Logs are clean.")